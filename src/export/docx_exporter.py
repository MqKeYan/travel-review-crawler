"""
模块名称：DOCX 格式导出器

功能说明：
    - 将评论数据导出为 Word DOCX 格式
    - 带格式表格展示
    - 适合生成阅读性强的评论汇总报告

依赖模块：
    - python-docx (第三方)
"""

import os
import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from src.export.base import BaseExporter
from src.models.review import STANDARD_FIELDS
from src.utils.exceptions import ExportError


class DocxExporter(BaseExporter):
    """Word DOCX 格式导出器"""

    format_name: str = "docx"
    file_extension: str = "docx"

    def export(
        self,
        reviews: list[dict],
        filepath: str,
        fields: list[str] | None = None,
    ) -> str:
        """
        导出为 DOCX 格式。

        生成包含标题、统计信息和数据表格的 Word 文档。

        Args:
            reviews: 评论数据列表
            filepath: 输出路径（不含扩展名）
            fields: 需要导出的字段列，None 则导出所有标准字段

        Returns:
            写入的文件完整路径

        Raises:
            ExportError: 写入文件失败
        """
        filepath = self._ensure_extension(filepath)

        # 确定字段和表头
        if fields:
            field_names = fields
            headers = self._get_headers(fields)
        else:
            field_names = [k for k, _ in STANDARD_FIELDS]
            headers = [v for _, v in STANDARD_FIELDS]

        try:
            doc = Document()

            # 设置默认字体为微软雅黑（支持中文）
            style = doc.styles["Normal"]
            font = style.font
            font.name = "微软雅黑"
            font.size = Pt(10)
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # ---- 文档标题 ----
            title = doc.add_heading("旅游评论数据导出报告", level=0)
            title.alignment = 1  # 居中

            # ---- 统计信息 ----
            doc.add_paragraph(f"导出时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"评论总数: {len(reviews)} 条")

            # ---- 创建数据表格 ----
            table = doc.add_table(rows=len(reviews) + 1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头行
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header

            # 数据行
            image_fields = {"image_urls", "avatar_url"}

            for row_idx, review in enumerate(reviews):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, field in enumerate(field_names):
                    value = review.get(field, "")

                    if field in image_fields and value:
                        # 图片/头像列：内嵌图片到单元格
                        paths = value if isinstance(value, list) else [value]
                        self._write_image_cell(row_cells[col_idx], paths, doc)
                    elif isinstance(value, list):
                        value = "; ".join(str(v) for v in value)
                        row_cells[col_idx].text = str(value)
                    else:
                        row_cells[col_idx].text = str(value)

            # 设置表格列宽（平均分配）
            page_width = doc.sections[0].page_width
            margin = doc.sections[0].left_margin + doc.sections[0].right_margin
            col_width = (page_width - margin) / len(headers)
            for row in table.rows:
                for cell in row.cells:
                    cell.width = col_width

            doc.save(filepath)

        except OSError as e:
            raise ExportError(f"DOCX 文件写入失败: {e}") from e
        except Exception as e:
            raise ExportError(f"DOCX 导出异常: {e}") from e

        return filepath

    def _get_headers(self, fields: list[str]) -> list[str]:
        """字段名 → 中文表头"""
        field_to_header = {k: v for k, v in STANDARD_FIELDS}
        return [field_to_header.get(f, f) for f in fields]

    def _write_image_cell(self, cell, image_paths: list[str], doc) -> None:
        """
        在表格单元格中嵌入图片。

        对于本地存在的图片文件，逐个嵌入到段落中；
        对于 URL 链接（下载失败的回退），显示为文本链接。

        Args:
            cell: docx 表格单元格对象
            image_paths: 图片路径列表（本地路径或 URL）
            doc: Document 对象
        """
        # 清空默认的空段落
        for p in cell.paragraphs:
            p.clear()

        for path in image_paths:
            if not isinstance(path, str) or not path:
                continue
            if path.startswith("http"):
                # URL 链接不写入，跳过
                continue
            if os.path.isfile(path):
                try:
                    self._embed_image_in_paragraph(cell.add_paragraph(), path)
                except Exception:
                    pass

    def _embed_image_in_paragraph(self, paragraph, image_path: str) -> None:
        """
        在段落中嵌入单张图片，自动缩放至合适尺寸。

        Args:
            paragraph: docx 段落对象
            image_path: 本地图片文件路径
        """
        from docx.shared import Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        max_width = Cm(5.5)     # 最大宽度 5.5cm
        max_height = Cm(4.5)    # 最大高度 4.5cm

        try:
            run = paragraph.add_run()
            run.add_picture(image_path, width=max_width, height=max_height)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            # 图片损坏或格式不支持，显示提示
            para_text = paragraph.add_run(f"[不支持格式] {os.path.basename(image_path)}")
            para_text.font.size = Pt(8)
